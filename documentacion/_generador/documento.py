# -*- coding: utf-8 -*-
"""Genera OTECH-GCS-SW-001 en .docx y lo convierte a PDF con Word.

    python documento.py

Salida en docs/. El PDF se produce automatizando Word (ExportAsFixedFormat) en
lugar de con un motor propio, porque así el índice, los campos de página y la
paginación del PDF coinciden exactamente con los del .docx que se entrega.
"""

import datetime
import os
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Emu

import contenido_req as R

AQUI = os.path.dirname(os.path.abspath(__file__))
DOCS = os.path.dirname(AQUI)
FIGS = os.path.join(DOCS, "figuras")
RAIZ = os.path.dirname(DOCS)

CODIGO = "OTECH-GCS-SW-001"
TITULO = "Plan de Aseguramiento y Especificación de Software"
SUBTITULO = ("Estación de Control Terrestre para Sistemas de Aeronave "
             "Pilotada a Distancia (RPAS / UAS)")
PRODUCTO = "OTECH-GroundStation"
REVISION = "A"
FECHA = datetime.date.today().strftime("%d/%m/%Y")
CLASIFICACION = "Uso interno / Presentación ante autoridad aeronáutica"

NAVY = RGBColor(0x0F, 0x1B, 0x3D)
CYAN = RGBColor(0x1E, 0x98, 0xC1)
GRIS = RGBColor(0x5A, 0x64, 0x72)
NAVY_HEX = "0F1B3D"
CYAN_HEX = "1E98C1"
FILA_HEX = "EEF2F7"
FUENTE = "Arial"


# ---------------------------------------------------------------------------
# Utilidades OOXML que python-docx no expone
# ---------------------------------------------------------------------------
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def sombrear(celda, hex_color):
    celda._tc.get_or_add_tcPr().append(_el("w:shd", **{"w:val": "clear", "w:fill": hex_color}))


def bordes_tabla(tabla, color="C4CCD8", size="4"):
    props = tabla._tbl.tblPr
    borders = _el("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{lado}", **{"w:val": "single", "w:sz": size,
                                           "w:space": "0", "w:color": color}))
    props.append(borders)


def repetir_encabezado(fila):
    fila._tr.get_or_add_trPr().append(_el("w:tblHeader", **{"w:val": "true"}))


def campo(parrafo, instruccion, texto_provisional="1"):
    """Inserta un campo de Word (PAGE, NUMPAGES, SEQ, TOC...)."""
    r1 = parrafo.add_run()
    r1._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    r2 = parrafo.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruccion
    r2._r.append(instr)
    r3 = parrafo.add_run()
    r3._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    r4 = parrafo.add_run(texto_provisional)
    r5 = parrafo.add_run()
    r5._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))
    return [r1, r2, r3, r4, r5]


def sin_espacio_despues(parrafo):
    parrafo.paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Constructor del documento
# ---------------------------------------------------------------------------
class Doc:
    def __init__(self):
        self.d = Document()
        self._estilos()
        self._pagina()
        self.n_fig = 0
        self.n_tab = 0

    # -- configuración base ------------------------------------------------
    def _estilos(self):
        n = self.d.styles["Normal"]
        n.font.name = FUENTE
        n.font.size = Pt(10)
        n.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.12
        n.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        n.element.rPr.rFonts.set(qn("w:eastAsia"), FUENTE)

        for nivel, (tam, color, antes) in enumerate(
                [(16, NAVY, 18), (12.5, NAVY, 14), (11, CYAN, 12), (10, GRIS, 10)], start=1):
            s = self.d.styles[f"Heading {nivel}"]
            s.font.name = FUENTE
            s.font.size = Pt(tam)
            s.font.bold = True
            s.font.color.rgb = color
            s.paragraph_format.space_before = Pt(antes)
            s.paragraph_format.space_after = Pt(6)
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            s.paragraph_format.keep_with_next = True

    def _pagina(self):
        s = self.d.sections[0]
        s.page_width, s.page_height = Cm(21.59), Cm(27.94)   # Carta
        s.left_margin = s.right_margin = Cm(2.3)
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.0)
        s.header_distance = Cm(1.1)
        s.footer_distance = Cm(1.0)

    @staticmethod
    def _tabuladores(parrafo, ancho_util_cm=16.99):
        """Tabulador central y derecho al ancho real de la caja de texto.

        Los del estilo Header/Footer heredado estan calculados para otro tamano de
        pagina y otros margenes: sin redefinirlos, el \\t del encabezado cae en una
        posicion ya rebasada y los dos textos salen pegados.
        """
        tabs = parrafo.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(ancho_util_cm / 2), WD_TAB_ALIGNMENT.CENTER)
        tabs.add_tab_stop(Cm(ancho_util_cm), WD_TAB_ALIGNMENT.RIGHT)

    def _encabezado_pie(self, seccion):
        seccion.different_first_page_header_footer = True

        enc = seccion.header.paragraphs[0]
        enc.text = ""
        enc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._tabuladores(enc)
        t = enc.add_run(f"{CODIGO}  ·  Rev. {REVISION}")
        t.font.size = Pt(8)
        t.font.color.rgb = GRIS
        t.font.name = FUENTE
        t2 = enc.add_run(f"\t\t{PRODUCTO}")
        t2.font.size = Pt(8)
        t2.font.color.rgb = GRIS
        t2.font.name = FUENTE
        self._linea_inferior(enc)

        pie = seccion.footer.paragraphs[0]
        pie.text = ""
        pie.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._tabuladores(pie)
        a = pie.add_run(CLASIFICACION)
        a.font.size = Pt(7.5)
        a.font.color.rgb = GRIS
        a.font.name = FUENTE
        b = pie.add_run("\t\tPágina ")
        b.font.size = Pt(8)
        b.font.color.rgb = GRIS
        b.font.name = FUENTE
        for r in campo(pie, "PAGE") + [pie.add_run(" de ")] + campo(pie, "NUMPAGES"):
            r.font.size = Pt(8)
            r.font.color.rgb = GRIS
            r.font.name = FUENTE

    @staticmethod
    def _linea_inferior(parrafo):
        pPr = parrafo._p.get_or_add_pPr()
        bdr = _el("w:pBdr")
        bdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "6",
                                      "w:space": "2", "w:color": "C4CCD8"}))
        pPr.append(bdr)

    # -- bloques de contenido ---------------------------------------------
    def h(self, nivel, texto):
        p = self.d.add_heading(texto, level=nivel)
        return p

    def p(self, texto, cursiva=False, tam=None, color=None, alinear=None, espacio=None):
        par = self.d.add_paragraph()
        r = par.add_run(texto)
        r.italic = cursiva
        r.font.name = FUENTE
        if tam:
            r.font.size = Pt(tam)
        if color:
            r.font.color.rgb = color
        if alinear is not None:
            par.alignment = alinear
        if espacio is not None:
            par.paragraph_format.space_after = Pt(espacio)
        return par

    def lista(self, elementos, vinetas=True):
        for e in elementos:
            par = self.d.add_paragraph(style="List Bullet" if vinetas else "List Number")
            par.paragraph_format.space_after = Pt(3)
            par.paragraph_format.left_indent = Cm(0.8)
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if isinstance(e, tuple):
                r = par.add_run(e[0] + "  ")
                r.bold = True
                r.font.name = FUENTE
                r.font.size = Pt(10)
                r2 = par.add_run(e[1])
                r2.font.name = FUENTE
                r2.font.size = Pt(10)
            else:
                r = par.add_run(e)
                r.font.name = FUENTE
                r.font.size = Pt(10)

    def nota(self, texto):
        t = self.d.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0)
        sombrear(c, "FFF6E0")
        bordes_tabla(t, color="D9B441", size="6")
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(texto)
        r.font.size = Pt(9)
        r.font.name = FUENTE
        r.font.color.rgb = RGBColor(0x5A, 0x42, 0x00)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.d.add_paragraph().paragraph_format.space_after = Pt(2)

    def figura(self, archivo, pie, ancho_cm=16.4):
        self.n_fig += 1
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.add_run().add_picture(os.path.join(FIGS, archivo), width=Cm(ancho_cm))

        cap = self.d.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        cap.style = self.d.styles["Caption"] if "Caption" in [s.name for s in self.d.styles] else cap.style
        r = cap.add_run("Figura ")
        for run in [r] + campo(cap, 'SEQ Figura \\* ARABIC', str(self.n_fig)):
            run.font.size = Pt(8.5)
            run.font.name = FUENTE
            run.font.color.rgb = GRIS
            run.bold = True
        r2 = cap.add_run(". " + pie)
        r2.font.size = Pt(8.5)
        r2.font.name = FUENTE
        r2.font.color.rgb = GRIS

    def tabla(self, cabeceras, filas, anchos=None, pie=None, tam=8.5, cabecera_tam=8.5):
        if pie:
            self.n_tab += 1
            cap = self.d.add_paragraph()
            cap.paragraph_format.space_before = Pt(10)
            cap.paragraph_format.space_after = Pt(3)
            r = cap.add_run("Tabla ")
            for run in [r] + campo(cap, 'SEQ Tabla \\* ARABIC', str(self.n_tab)):
                run.font.size = Pt(8.5)
                run.font.name = FUENTE
                run.font.color.rgb = GRIS
                run.bold = True
            r2 = cap.add_run(". " + pie)
            r2.font.size = Pt(8.5)
            r2.font.name = FUENTE
            r2.font.color.rgb = GRIS

        t = self.d.add_table(rows=1, cols=len(cabeceras))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        bordes_tabla(t)

        hdr = t.rows[0]
        repetir_encabezado(hdr)
        for i, texto in enumerate(cabeceras):
            c = hdr.cells[i]
            sombrear(c, NAVY_HEX)
            par = c.paragraphs[0]
            par.text = ""
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.space_before = Pt(2)
            # El estilo Normal justifica, y en una celda estrecha eso reparte un
            # encabezado de dos palabras a lo ancho con huecos enormes.
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = par.add_run(texto)
            run.bold = True
            run.font.size = Pt(cabecera_tam)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = FUENTE

        for n, fila in enumerate(filas):
            celdas = t.add_row().cells
            for i, valor in enumerate(fila):
                c = celdas[i]
                if n % 2 == 1:
                    sombrear(c, FILA_HEX)
                par = c.paragraphs[0]
                par.text = ""
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.space_before = Pt(2)
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = par.add_run(str(valor))
                run.font.size = Pt(tam)
                run.font.name = FUENTE

        if anchos:
            total = sum(anchos)
            util = 16.99  # cm útiles con márgenes de 2,3 cm
            for fila in t.rows:
                for i, c in enumerate(fila.cells):
                    c.width = Cm(util * anchos[i] / total)

        self.d.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def salto(self):
        self.d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Contenido
# ---------------------------------------------------------------------------
def portada(doc):
    d = doc.d
    logo = os.path.join(RAIZ, "logo.png")

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(6)
    if os.path.exists(logo):
        p.add_run().add_picture(logo, width=Cm(5.2))

    for texto, tam, color, negrita, antes, despues in [
        ("OTECH", 13, CYAN, True, 10, 30),
        (TITULO.upper(), 21, NAVY, True, 0, 4),
        (SUBTITULO, 12.5, GRIS, False, 0, 26),
        (PRODUCTO, 17, NAVY, True, 0, 2),
        ("Versión del producto 5.0.8  ·  Compilación de referencia e0816c9", 9.5, GRIS, False, 0, 34),
    ]:
        par = d.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(antes)
        par.paragraph_format.space_after = Pt(despues)
        r = par.add_run(texto)
        r.font.size = Pt(tam)
        r.font.color.rgb = color
        r.bold = negrita
        r.font.name = FUENTE

    doc.tabla(
        ["Campo", "Contenido"],
        [
            ("Código de documento", CODIGO),
            ("Revisión", f"Rev. {REVISION}"),
            ("Fecha de emisión", FECHA),
            ("Estado del documento", "Borrador para revisión — no constituye una declaración de conformidad"),
            ("Organización responsable", "OTECH"),
            ("Elemento de configuración", f"{PRODUCTO} v5.0.8 (rama otech-branding)"),
            ("Clasificación", CLASIFICACION),
            ("Idioma", "Español (México)"),
        ],
        anchos=[4, 12], tam=9.5, cabecera_tam=9.5)

    doc.p("El presente documento se emite para su revisión técnica. Las declaraciones de "
          "cumplimiento normativo contenidas en el capítulo 3 y en las matrices del capítulo 7 "
          "están sujetas a la verificación independiente descrita en el capítulo 9 y al estado de "
          "madurez declarado en el capítulo 13.",
          cursiva=True, tam=8.5, color=GRIS)
    doc.salto()


def control(doc):
    doc.h(1, "Hoja de control del documento")

    doc.h(2, "Historial de revisiones")
    doc.tabla(["Rev.", "Fecha", "Autor", "Descripción del cambio"],
              [("A", FECHA, "Ingeniería de Software OTECH",
                "Emisión inicial. Establece el marco normativo, la línea base de requisitos de alto "
                "y bajo nivel, la matriz de trazabilidad y el ciclo de vida de desarrollo.")],
              anchos=[1.2, 2, 3.5, 9.3], tam=9)

    doc.h(2, "Elaboración, revisión y aprobación")
    doc.tabla(["Función", "Nombre y cargo", "Firma", "Fecha"],
              [("Elabora", "Ingeniería de Software", "", ""),
               ("Revisa", "Aseguramiento de Calidad de Software", "", ""),
               ("Revisa", "Seguridad Operacional (SMS)", "", ""),
               ("Aprueba", "Dirección Técnica", "", ""),
               ("Acepta", "Explotador / Responsable de la operación", "", "")],
              anchos=[2.5, 6.5, 4, 3], tam=9)

    doc.h(2, "Distribución")
    doc.tabla(["Destinatario", "Propósito"],
              [("Autoridad aeronáutica competente (AFAC)",
                "Sustento técnico del sistema de estación terrestre asociado a la solicitud de "
                "autorización de operación."),
               ("Explotador / operador RPAS",
                "Definición de las capacidades y limitaciones del software empleado como estación "
                "de pilotaje a distancia."),
               ("Ingeniería y calidad OTECH",
                "Línea base de requisitos y de verificación para el desarrollo y el mantenimiento."),
               ("Piloto a distancia",
                "Referencia de las funciones de seguridad y de las limitaciones declaradas.")],
              anchos=[5, 11], tam=9)

    doc.h(2, "Control de cambios")
    doc.p("Toda modificación de este documento requiere la reemisión completa con incremento de "
          "revisión, la actualización del historial anterior y una nueva ronda de aprobación. Los "
          "cambios que afecten a un requisito ya verificado obligan a reejecutar los casos de "
          "verificación asociados según la matriz de trazabilidad del capítulo 7. No se admiten "
          "modificaciones parciales ni anotaciones manuscritas sobre ejemplares emitidos.")
    doc.salto()


def indices(doc):
    d = doc.d
    doc.h(1, "Índice")
    p = d.add_paragraph()
    campo(p, 'TOC \\o "1-3" \\h \\z \\u',
          "Actualice este índice en Word: Referencias > Actualizar tabla > Actualizar toda la tabla.")
    doc.salto()

    doc.h(1, "Índice de figuras")
    p = d.add_paragraph()
    campo(p, 'TOC \\h \\z \\c "Figura"', "Índice de figuras (actualizar campo).")

    doc.h(1, "Índice de tablas")
    p = d.add_paragraph()
    campo(p, 'TOC \\h \\z \\c "Tabla"', "Índice de tablas (actualizar campo).")
    doc.salto()


def cap1(doc):
    doc.h(1, "1. Introducción")

    doc.h(2, "1.1 Objeto")
    doc.p(f"Este documento especifica los requisitos, el ciclo de vida de desarrollo y la estrategia "
          f"de verificación del software {PRODUCTO}, empleado como estación de control terrestre "
          f"para la operación de sistemas de aeronave pilotada a distancia. Su objeto es proporcionar "
          f"a la autoridad aeronáutica competente y al explotador la evidencia estructurada necesaria "
          f"para valorar la aptitud del software dentro del sistema RPAS al que da servicio.")
    doc.p("El documento cumple una triple función: es la especificación técnica contra la que se "
          "desarrolla el software, es el plan de aseguramiento que define cómo se demuestra dicha "
          "especificación, y es el registro de trazabilidad que vincula cada exigencia normativa con "
          "una función implementada y con una evidencia de verificación.")

    doc.h(2, "1.2 Alcance")
    doc.p("El alcance comprende la totalidad del software de la estación de control terrestre: la "
          "interfaz de pilotaje, la gestión del enlace de mando y control, la planificación y "
          "ejecución de la misión, las funciones de contención y seguridad operacional, la gestión "
          "de la carga útil, la identificación remota y el registro de la operación.")
    doc.p("Quedan expresamente fuera del alcance de este documento:")
    doc.lista([
        "El software embarcado en la aeronave (controlador de vuelo y sus periféricos), objeto de su "
        "propia documentación de aeronavegabilidad.",
        "El equipo físico sobre el que se ejecuta la estación (computador, radio control, radioenlace), "
        "cuya calificación ambiental y eléctrica se documenta por separado.",
        "Los procedimientos operacionales del explotador, el manual de operaciones y la formación "
        "del piloto a distancia.",
        "La evaluación de riesgo específica de cada operación, que corresponde al explotador y se "
        "documenta mediante la metodología aceptada por la autoridad.",
    ])
    doc.p("El software es un medio para ejecutar la operación, no un sustituto de la evaluación de "
          "riesgo ni de los procedimientos del explotador. Las declaraciones de este documento "
          "aplican al software y a sus interfaces, no a la operación en su conjunto.")

    doc.h(2, "1.3 Identificación del elemento de configuración")
    doc.tabla(["Atributo", "Valor"],
              [("Denominación del producto", PRODUCTO),
               ("Versión del producto", "5.0.8"),
               ("Línea base de código", "Rama otech-branding, commit de referencia e0816c9"),
               ("Software preexistente de origen", "QGroundControl v5.0.8 (véase 2.7)"),
               ("Identificador de aplicación (escritorio)", "com.otech.groundstation"),
               ("Identificador de paquete (Android)", "org.mavlink.qgroundcontrol"),
               ("Versión del esquema de configuración", "9"),
               ("Lenguajes de implementación", "C++17 y QML"),
               ("Volumen del código fuente propio y adaptado",
                "328 unidades de traducción C++, 384 cabeceras, 422 módulos QML"),
               ("Protocolo de interfaz con la aeronave", "MAVLink 2.0 (biblioteca c_library_v2)")],
              anchos=[5.5, 10.5], tam=9,
              pie="Identificación del elemento de configuración objeto de este documento.")
    doc.nota("El identificador de paquete de Android conserva el valor heredado del software "
             "preexistente y no coincide con el identificador de aplicación de escritorio. Se "
             "documenta aquí de forma explícita porque es el valor por el que la autoridad o el "
             "explotador identificarán la instalación en el dispositivo. Su armonización está "
             "recogida en el plan de acción del capítulo 14.")

    doc.h(2, "1.4 Destinatarios")
    doc.lista([
        ("Autoridad aeronáutica.", "Valoración de la aptitud del software como parte del sistema RPAS."),
        ("Explotador.", "Conocimiento de las capacidades, limitaciones y supuestos del software."),
        ("Ingeniería de software.", "Línea base de requisitos para el desarrollo y el mantenimiento."),
        ("Aseguramiento de calidad.", "Base para la planificación y la auditoría de la verificación."),
    ])

    doc.h(2, "1.5 Documentos aplicables")
    doc.p("Los documentos aplicables establecen exigencias de obligado cumplimiento para el software "
          "objeto de este documento.")
    doc.tabla(["Ref.", "Documento", "Aplicación"],
              [("[A-1]", "NOM-107-SCT3-2019. Requerimientos para operar un sistema de aeronave "
                         "pilotada a distancia (RPAS) en el espacio aéreo mexicano.",
                "Marco nacional de operación. Origen de los requisitos de operación, contención, "
                "información al piloto y registro."),
               ("[A-2]", "Ley de Aviación Civil y su Reglamento (Estados Unidos Mexicanos).",
                "Marco legal de referencia para la operación en espacio aéreo nacional."),
               ("[A-3]", "OACI, Anexo 6 Parte IV. Operación de aeronaves — Sistemas de aeronaves "
                         "pilotadas a distancia.",
                "Exigencias de la estación de pilotaje a distancia."),
               ("[A-4]", "OACI, Anexo 2. Reglamento del aire.",
                "Reglas aplicables al tránsito y a la separación."),
               ("[A-5]", "OACI, Anexo 13. Investigación de accidentes e incidentes de aviación.",
                "Exigencias sobre conservación y disponibilidad de la evidencia registrada."),
               ("[A-6]", "ASTM F3411-22a. Standard Specification for Remote ID and Tracking.",
                "Formato y contenido de la identificación remota."),
               ("[A-7]", "RTCA DO-278A / EUROCAE ED-109A. Software Integrity Assurance "
                         "Considerations for CNS/ATM Systems.",
                "Norma de aseguramiento adoptada para el software de tierra (véase 3.4).")],
              anchos=[1.2, 8.3, 6.5], tam=8.5,
              pie="Documentos aplicables.")

    doc.h(2, "1.6 Documentos de referencia")
    doc.tabla(["Ref.", "Documento", "Uso"],
              [("[R-1]", "OACI Doc 10019 AN/507. Manual sobre sistemas de aeronaves pilotadas a distancia.",
                "Guía de interpretación de las funciones de la estación de pilotaje."),
               ("[R-2]", "OACI Doc 9859. Manual de gestión de la seguridad operacional.",
                "Metodología de gestión del riesgo y de clasificación de severidad."),
               ("[R-3]", "JARUS, Specific Operations Risk Assessment (SORA), Objetivos de Seguridad "
                         "Operacional OSO#01 a OSO#24.",
                "Metodología de referencia para derivar exigencias de robustez del software."),
               ("[R-4]", "Reglamento de Ejecución (UE) 2019/947 y Reglamento Delegado (UE) 2019/945.",
                "Referencia comparada para la categoría de operación específica."),
               ("[R-5]", "RTCA DO-178C / EUROCAE ED-12C. Software Considerations in Airborne Systems.",
                "Referencia de objetivos de proceso adaptados al software de tierra."),
               ("[R-6]", "RTCA DO-326A / EUROCAE ED-202A. Airworthiness Security Process Specification.",
                "Base de los requisitos de seguridad de la información del capítulo 12."),
               ("[R-7]", "SAE ARP4754A y ARP4761.",
                "Desarrollo de sistemas y metodología de evaluación de seguridad."),
               ("[R-8]", "SAE ARP4102/4 y ARP4102/7.",
                "Convenciones de alertas y de presentación en cabina."),
               ("[R-9]", "ISO/IEC/IEEE 29148:2018 e ISO/IEC/IEEE 12207:2017.",
                "Ingeniería de requisitos y procesos del ciclo de vida del software."),
               ("[R-10]", "ASTM F3548-21. UAS Traffic Management (UTM) UAS Service Supplier "
                          "Interoperability.",
                "Interfaz con proveedores de servicios UTM."),
               ("[R-11]", "ASTM F3201-16. Ensuring Dependability of Software Used in Unmanned "
                          "Aircraft Systems.",
                "Buenas prácticas de dependabilidad aplicables a software de UAS.")],
              anchos=[1.2, 8.3, 6.5], tam=8.5,
              pie="Documentos de referencia.")

    doc.h(2, "1.7 Estructura del documento")
    doc.p("Los capítulos 2 y 3 establecen qué es el sistema y contra qué normas se evalúa. El "
          "capítulo 4 identifica qué puede fallar y con qué consecuencia, del que se derivan "
          "requisitos adicionales. Los capítulos 5 y 6 contienen la especificación propiamente "
          "dicha, en dos niveles de detalle. El capítulo 7 demuestra que ambas especificaciones "
          "cubren la normativa y se verifican. Los capítulos 8 y 9 describen cómo se desarrolla y "
          "cómo se comprueba. Los capítulos 10 a 12 cubren los procesos de soporte. El capítulo 13 "
          "declara sin atenuantes el estado real de madurez y el 14 fija el plan de acción "
          "pendiente.")
    doc.salto()


def cap2(doc):
    doc.h(1, "2. Descripción del sistema")

    doc.h(2, "2.1 Concepto de operación")
    doc.p("El software constituye el puesto de pilotaje a distancia desde el que un piloto "
          "cualificado planifica, supervisa y manda una o varias aeronaves no tripuladas. El piloto "
          "conserva en todo momento la responsabilidad de la operación; el software es el medio a "
          "través del cual ejerce esa responsabilidad, y está diseñado bajo la premisa de que la "
          "aeronave mantiene sus propios mecanismos autónomos de protección con independencia del "
          "estado de la estación.")
    doc.p("El perfil de operación previsto comprende:")
    doc.lista([
        ("Operación en alcance visual (VLOS).", "El piloto mantiene contacto visual directo con la "
         "aeronave y emplea la estación como fuente de información complementaria y como medio de mando."),
        ("Operación en alcance visual ampliado (EVLOS) y más allá del alcance visual (BVLOS).",
         "La estación pasa a ser la fuente primaria de conciencia situacional, lo que eleva las "
         "exigencias de integridad y disponibilidad tratadas en el capítulo 3."),
        ("Operación desde puesto fijo o desde radio control portatil.", "El software se ejecuta "
         "indistintamente sobre un computador de escritorio o sobre un radio control con sistema "
         "operativo Android, con idéntica funcionalidad operacional."),
    ])
    doc.p("Escenario nominal: el piloto prepara el plan de vuelo y las geocercas antes del vuelo, "
          "verifica las precondiciones comunicadas por la aeronave, transfiere el plan, arma y "
          "despega bajo confirmación explícita, supervisa la ejecución mediante telemetría y video, "
          "y dispone en todo momento del mando de retorno al punto de lanzamiento. Toda la sesión "
          "queda registrada para su análisis posterior.")

    doc.h(2, "2.2 Frontera del sistema y supuestos")
    doc.p("La frontera del sistema se sitúa en las interfaces enumeradas en 2.4. Todo lo que queda "
          "al otro lado de esas interfaces se considera externo y se rige por los siguientes "
          "supuestos, que condicionan la validez de los requisitos especificados:")
    doc.lista([
        ("S-1.", "La aeronave implementa mecanismos autónomos de protección ante pérdida de enlace, "
                 "batería baja y vulneración de geocerca, independientes de la estación."),
        ("S-2.", "La aeronave es la autoridad última sobre su propia contención; la validación que "
                 "realiza la estación es una barrera adicional, no la única."),
        ("S-3.", "El enlace de radio entre la estación y la aeronave se dimensiona y se verifica "
                 "como parte del sistema RPAS, no del software."),
        ("S-4.", "El piloto a distancia posee la habilitación exigida y conoce los procedimientos "
                 "de contingencia del explotador."),
        ("S-5.", "El equipo anfitrión cumple las características mínimas declaradas en 2.5 y su "
                 "reloj esta sincronizado con una fuente horaria fiable."),
        ("S-6.", "La cartografía y los datos de elevación empleados proceden de fuentes declaradas "
                 "y se han descargado antes de la operación cuando esta se realice sin cobertura."),
    ])
    doc.nota("Si alguno de estos supuestos no se cumple en una operación concreta, las conclusiones "
             "del análisis de seguridad del capítulo 4 dejan de ser válidas para esa operación y el "
             "explotador debe reevaluar el riesgo.")

    doc.h(2, "2.3 Arquitectura lógica")
    doc.p("El software se organiza en capas con dependencia descendente estricta: ninguna capa "
          "invoca servicios de una capa superior. Esta separación es la que permite asignar cada "
          "requisito de bajo nivel a un componente identificable y acotar el alcance de la "
          "verificación cuando se modifica una función.")
    doc.figura("fig02_arquitectura.png",
               "Arquitectura lógica en capas y frontera del sistema.")
    doc.p("La capa de comunicaciones aísla por completo el resto del software del medio físico "
          "empleado: la capa de aplicación opera del mismo modo tanto si la aeronave está conectada "
          "por radio serie como por red o por un registro reproducido. Esa propiedad es la que hace "
          "posible verificar la mayor parte de los requisitos funcionales en banco, sin aeronave "
          "real, mediante el enlace simulado.")

    doc.h(2, "2.4 Interfaces externas")
    doc.tabla(["Ref.", "Interfaz", "Dirección", "Protocolo / medio", "Criticidad"],
              [("IF-01", "Aeronave (enlace de mando y control)", "Bidireccional",
                "MAVLink 2.0 sobre serie, UDP, TCP o Bluetooth", "Crítica"),
               ("IF-02", "Receptor GNSS / correcciones RTK", "Entrada",
                "NMEA y RTCM3 sobre serie o red", "Esencial"),
               ("IF-03", "Servidor de cartografía y elevación", "Entrada",
                "HTTPS, con almacén local para operación sin conexión", "Esencial"),
               ("IF-04", "Fuente de video de la carga útil", "Entrada",
                "RTSP, UDP (H.264/H.265), TCP-MPEG2TS, UVC local", "Esencial"),
               ("IF-05", "Proveedor de servicios UTM", "Bidireccional",
                "API REST sobre HTTPS", "Esencial"),
               ("IF-06", "Receptor ADS-B", "Entrada", "TCP (SBS-1 / Beast)", "Esencial"),
               ("IF-07", "Mando manual del piloto", "Entrada",
                "HID (joystick o palanca de mando)", "Esencial"),
               ("IF-08", "Almacenamiento local de registros", "Salida",
                "Sistema de archivos del equipo anfitrión", "Crítica")],
              anchos=[1.2, 5, 2.2, 5.6, 2], tam=8.5,
              pie="Interfaces externas del software y su criticidad.")

    doc.h(2, "2.5 Entorno de ejecución")
    doc.tabla(["Plataforma", "Configuración soportada", "Equipo de referencia para verificación"],
              [("Windows", "Windows 10 y 11, 64 bits (x86-64)",
                "Procesador de 4 núcleos, 16 GB de memoria, GPU con soporte OpenGL 3.3"),
               ("Android", "Android 9 (API 28) o superior, arquitectura arm64-v8a",
                "Radio control SIYI, 8 GB de memoria, pantalla de 7 pulgadas")],
              anchos=[3, 6.5, 6.5], tam=9,
              pie="Plataformas soportadas y equipos de referencia.")
    doc.p("Los umbrales cuantitativos de los requisitos de desempeño del dominio PER se refieren "
          "exclusivamente a los equipos de referencia declarados en esta tabla. Su verificación "
          "sobre un equipo distinto no es extrapolable.")

    doc.h(2, "2.6 Entorno de desarrollo y cadena de herramientas")
    doc.tabla(["Elemento", "Versión y configuración"],
              [("Marco de aplicación", "Qt 6.8.3 (Core, Quick, Positioning, Location, Multimedia, "
                                       "SerialPort, Bluetooth, Charts)"),
               ("Compilador (Windows)", "MSVC 2022, conjunto de herramientas x64"),
               ("Compilador (Android)", "NDK r26b (26.1.10909125), ABI arm64-v8a"),
               ("Sistema de construcción", "CMake con generador Ninja Multi-Config"),
               ("Biblioteca de protocolo", "MAVLink c_library_v2, generada en la construcción"),
               ("Subsistema de video", "GStreamer 1.0"),
               ("Control de versiones", "Git; la versión del producto se deriva de la etiqueta del "
                                        "repositorio en el momento de compilar"),
               ("Marco de pruebas", "Qt Test, integrado en el árbol test/ del repositorio")],
              anchos=[5, 11], tam=9,
              pie="Cadena de herramientas de desarrollo.")
    doc.p("Ninguna herramienta de la cadena elimina o sustituye una actividad de verificación, por "
          "lo que no se requiere cualificación de herramientas en el sentido de DO-330. Si en el "
          "futuro se automatiza la generación de código o se sustituye una revisión manual por el "
          "resultado de una herramienta, dicha herramienta deberá cualificarse antes de acreditar su "
          "resultado como evidencia.")

    doc.h(2, "2.7 Software preexistente y componentes de terceros")
    doc.p("El producto se construye sobre software preexistente de código abierto. Este hecho se "
          "declara de forma explícita porque condiciona la estrategia de aseguramiento: no es "
          "posible acreditar el cumplimiento de objetivos de proceso sobre código cuyo desarrollo "
          "original no fue gobernado por este plan. La estrategia adoptada, conforme al tratamiento "
          "de software previamente desarrollado de [A-7], consiste en acreditar el producto mediante "
          "verificación sobre el software integrado, y no mediante la reconstrucción retroactiva de "
          "la evidencia de proceso del componente de origen.")
    doc.tabla(["Componente", "Versión", "Licencia", "Función"],
              [("QGroundControl", "5.0.8", "Apache 2.0 / GPLv3",
                "Base funcional de la estación de control terrestre"),
               ("Qt", "6.8.3", "LGPLv3", "Marco de aplicación e interfaz de usuario"),
               ("GStreamer", "1.0", "LGPLv2.1", "Recepción, decodificación y grabación de video"),
               ("MAVLink c_library_v2", "Generada", "MIT", "Serialización del protocolo con la aeronave"),
               ("Qt Charts", "6.8.3", "GPLv3", "Representación gráfica en las vistas de análisis")],
              anchos=[4.5, 2.5, 3.5, 5.5], tam=9,
              pie="Componentes preexistentes y de terceros.")
    doc.p("Las modificaciones introducidas por OTECH sobre el software preexistente comprenden la "
          "identidad visual del producto, la consola de mensajes MAVLink de la vista de vuelo, la "
          "pantalla de presentación, la revisión de la presentación de instrumentos y gráficos, y "
          "cinco correcciones de defectos necesarias para la construcción y la estabilidad en las "
          "plataformas soportadas. Estas modificaciones sí están íntegramente cubiertas por los "
          "requisitos de bajo nivel del capítulo 6.")
    doc.salto()


def cap3(doc):
    doc.h(1, "3. Marco normativo y estrategia de cumplimiento")

    doc.h(2, "3.1 Marco nacional")
    doc.p("La operación de sistemas de aeronave pilotada a distancia en el espacio aéreo mexicano se "
          "rige por la NOM-107-SCT3-2019 [A-1], en el marco de la Ley de Aviación Civil y su "
          "Reglamento [A-2], bajo la competencia de la Agencia Federal de Aviación Civil. La norma "
          "regula la operación y no certifica productos de software de forma independiente; sus "
          "exigencias alcanzan al software en la medida en que este es el medio por el que se "
          "satisfacen obligaciones del operador.")
    doc.p("Este documento traduce esas obligaciones en requisitos verificables sobre el software. "
          "La correspondencia completa entre cada numeral de la norma y los requisitos que lo "
          "materializan se recoge en la tabla de trazabilidad normativa del apartado 7.2.")

    doc.h(2, "3.2 Marco internacional")
    doc.p("Se adopta como referencia el Anexo 6 Parte IV [A-3] y el Doc 10019 [R-1] de la OACI en "
          "lo relativo a las funciones exigibles a una estación de pilotaje a distancia, el Anexo 2 "
          "[A-4] en cuanto a las reglas del aire aplicables, y el Anexo 13 [A-5] en cuanto a la "
          "conservación y disponibilidad de la evidencia registrada para la investigación de sucesos.")

    doc.h(2, "3.3 Metodología de riesgo de referencia")
    doc.p("Para derivar el nivel de robustez exigible a cada función del software se emplea como "
          "referencia metodológica la evaluación SORA [R-3], ampliamente aceptada en el ámbito "
          "internacional para operaciones de la categoría específica. SORA no es de aplicación "
          "obligatoria en el marco nacional; se utiliza aquí como método estructurado para "
          "justificar decisiones de diseño, no como declaración de cumplimiento de un marco "
          "regulatorio extranjero.")
    doc.p("Los objetivos de seguridad operacional (OSO) con incidencia directa sobre el software de "
          "la estación son los siguientes:")
    doc.tabla(["OSO", "Objetivo", "Incidencia sobre el software"],
              [("OSO#05", "Diseño del UAS considerando la seguridad y la fiabilidad del sistema",
                "Degradación controlada, estabilidad prolongada y ausencia de modos de fallo únicos "
                "en la cadena de mando."),
               ("OSO#06", "Desempeño del enlace C3",
                "Detección y notificación de la pérdida de enlace, redundancia de enlaces e "
                "indicadores de calidad."),
               ("OSO#08", "Procedimientos operacionales definidos y validados",
                "Gestión verificable de parámetros y cotejo del plan residente en la aeronave."),
               ("OSO#10", "Contención del volumen de operación",
                "Definición, transferencia y validación de geocercas y de puntos de recuperación."),
               ("OSO#13", "Sistemas externos que soportan la operación",
                "Estado del subsistema de identificación remota e integridad de los enlaces con "
                "servicios externos."),
               ("OSO#16", "Procedimientos multitripulación",
                "Segregación de la información y del mando en operación multiaeronave."),
               ("OSO#19", "Recuperación ante error humano",
                "Confirmación de mandos críticos y presentación inequívoca del estado."),
               ("OSO#20", "Interfaz humano-máquina adecuada",
                "Prioridad de las alertas, no interferencia del video y legibilidad en exteriores.")],
              anchos=[1.6, 6, 8.4], tam=8.5,
              pie="Objetivos de seguridad operacional con incidencia sobre el software.")

    doc.h(2, "3.4 Norma de aseguramiento de software adoptada")
    doc.p("El software objeto de este documento no se ejecuta a bordo de la aeronave, por lo que "
          "DO-178C [R-5] no le es directamente aplicable. La norma adoptada es DO-278A / ED-109A "
          "[A-7], concebida precisamente para software de sistemas de tierra, que conserva la "
          "estructura de objetivos de DO-178C y define seis niveles de aseguramiento (AL1 a AL6) "
          "asignados en función de la severidad de la condición de fallo a nivel de sistema.")

    doc.h(2, "3.5 Determinación del nivel de aseguramiento")
    doc.p("La asignación del nivel de aseguramiento parte del análisis funcional de peligros del "
          "capítulo 4. La conclusión de dicho análisis es que ninguna condición de fallo del "
          "software de la estación alcanza por sí sola la categoría de catastrófica, en virtud del "
          "supuesto S-1: la aeronave conserva mecanismos autónomos de protección que actúan con "
          "independencia del estado de la estación. La condición de fallo más severa identificada, "
          "la pérdida total de la capacidad de mando (FC-01), se clasifica como mayor porque "
          "desemboca en la ejecución del procedimiento autónomo de la aeronave y no en una pérdida "
          "de control.")
    doc.tabla(["Severidad de la condición de fallo", "Nivel DO-278A", "Aplicabilidad al presente producto"],
              [("Catastrófica", "AL1", "No aplicable. Ninguna condición de fallo alcanza esta severidad "
                                       "bajo el supuesto S-1."),
               ("Peligrosa", "AL2", "No aplicable en el perfil de operación declarado en 2.1."),
               ("Mayor", "AL3 / AL4", "Aplicable. Nivel adoptado (véase la justificación siguiente)."),
               ("Menor", "AL5", "Aplicable a funciones sin incidencia en la seguridad, como el "
                                "intercambio de planes o la generación de patrones de barrido."),
               ("Sin efecto", "AL6", "Aplicable a funciones estrictamente estéticas o de comodidad.")],
              anchos=[4.5, 2.5, 9], tam=9,
              pie="Correspondencia entre severidad y nivel de aseguramiento DO-278A.")
    doc.p("Nivel adoptado: AL4 como línea base para las funciones clasificadas como críticas en los "
          "capítulos 5 y 6, en operaciones VLOS con la aeronave sobre zona no congestionada. Se "
          "adopta AL3 para esas mismas funciones cuando la operación se realice más allá del "
          "alcance visual o sobre concentraciones de personas, dado que en esos escenarios la "
          "estación pasa a ser la fuente primaria de conciencia situacional y desaparece la "
          "mitigación que aporta la observación directa del piloto.")
    doc.nota("La asignación de nivel de aseguramiento es una propuesta técnica del fabricante "
             "sustentada en el análisis del capítulo 4. Corresponde a la autoridad competente su "
             "aceptación, modificación o rechazo. El estado actual de cumplimiento de los objetivos "
             "asociados a AL4 se declara sin atenuantes en el capítulo 13.")

    doc.h(2, "3.6 Estrategia de cumplimiento")
    doc.p("La estrategia se articula en cinco líneas:")
    doc.lista([
        ("Trazabilidad completa.", "Toda exigencia normativa se traduce en al menos un requisito de "
         "alto nivel, este en uno o más requisitos de bajo nivel asignados a un componente "
         "identificable, y cada requisito de alto nivel en al menos un caso de verificación."),
        ("Verificación sobre el producto integrado.", "Dado el origen del software (2.7), la "
         "evidencia se construye verificando el comportamiento del producto, no reconstruyendo la "
         "evidencia de proceso del componente preexistente."),
        ("Independencia de la verificación.", "Los casos que verifican requisitos de funciones "
         "críticas los ejecuta personal distinto del que implementó la función."),
        ("Barreras redundantes.", "Ninguna función de seguridad depende exclusivamente del software "
         "de la estación; la contención efectiva reside en la aeronave y el software actúa como "
         "barrera adicional."),
        ("Declaración honesta del estado.", "El capítulo 13 declara que requisitos están verificados "
         "con evidencia registrada y cuáles no, sin presentar como acreditado lo que no lo está."),
    ])
    doc.salto()


def cap4(doc):
    doc.h(1, "4. Evaluación funcional de peligros")

    doc.h(2, "4.1 Metodología")
    doc.p("Se aplica un análisis funcional de peligros conforme a la orientación de ARP4761 [R-7], "
          "adaptado al alcance de un elemento de tierra. Para cada función del software se "
          "identifica su condición de fallo (pérdida de la función, función errónea no detectada y "
          "función no solicitada), se clasifica su severidad a nivel del sistema RPAS completo, se "
          "identifican las mitigaciones existentes y se derivan los requisitos necesarios para "
          "sostener esas mitigaciones.")
    doc.p("La clasificación se efectúa siempre a nivel de sistema, nunca a nivel de componente: la "
          "consecuencia relevante no es que un módulo de software falle, sino en qué situación queda "
          "la aeronave y el entorno sobrevolado cuando eso ocurre.")

    doc.h(2, "4.2 Clasificación de severidad")
    doc.tabla(["Severidad", "Definición aplicada en este análisis"],
              [("Catastrófica", "Pérdida de control de la aeronave con consecuencias sobre terceros "
                                "en tierra o sobre otra aeronave."),
               ("Peligrosa", "Reducción grave de los márgenes de seguridad; el piloto no puede "
                             "ejecutar el procedimiento de contingencia previsto."),
               ("Mayor", "Reducción significativa de los márgenes de seguridad o aumento apreciable "
                         "de la carga de trabajo del piloto; la operación debe interrumpirse."),
               ("Menor", "Molestia operativa o pérdida de una función accesoria sin efecto sobre la "
                         "seguridad de la operación en curso."),
               ("Sin efecto", "Sin consecuencia sobre la seguridad ni sobre la capacidad operativa.")],
              anchos=[3, 13], tam=9,
              pie="Escala de severidad empleada en el análisis funcional de peligros.")

    doc.h(2, "4.3 Condiciones de fallo identificadas")
    doc.tabla(["Ref.", "Condición de fallo", "Sev.", "Mitigación", "Requisitos asociados"],
              [(a, b, c, d, e) for a, b, c, d, e in R.FHA],
              anchos=[1, 4, 1.2, 5.8, 4], tam=8,
              pie="Condiciones de fallo, severidad y mitigaciones.")
    doc.p("Ninguna condición de fallo alcanza severidad peligrosa ni catastrófica. Esta conclusión "
          "depende íntegramente del supuesto S-1: si se opera una aeronave que carezca de "
          "mecanismos autónomos de protección ante pérdida de enlace, la condición FC-01 asciende "
          "a peligrosa y el nivel de aseguramiento exigible al software se eleva en consecuencia.")

    doc.h(2, "4.4 Requisitos derivados de seguridad")
    doc.p("Los siguientes requisitos no proceden de una exigencia normativa explícita, sino del "
          "propio análisis de peligros. Se identifican por separado porque, conforme a [A-7], los "
          "requisitos derivados deben remitirse al proceso de evaluación de seguridad del sistema "
          "para valorar su efecto:")
    doc.tabla(["Ref.", "Requisito derivado", "Requisitos de alto nivel que lo materializan"],
              [(a, b, c) for a, b, c in R.DERIVADOS],
              anchos=[1.2, 9.3, 5.5], tam=8.5,
              pie="Requisitos derivados del análisis de seguridad.")
    doc.salto()


def cap5(doc):
    doc.h(1, "5. Requisitos de alto nivel")

    doc.h(2, "5.1 Convención de identificación")
    doc.p("Cada requisito de alto nivel se identifica con la forma OTECH-HLR-DDD-nnn, donde DDD es "
          "el código de dominio funcional y nnn un número correlativo dentro del dominio. Los "
          "identificadores son permanentes: un requisito suprimido conserva su identificador marcado "
          "como obsoleto y su número no se reutiliza, de modo que las referencias de documentos "
          "anteriores siguen resolviéndose. En las tablas siguientes se omite el prefijo OTECH- por "
          "economía de espacio.")
    doc.tabla(["Código", "Dominio funcional"], [(c, n) for c, n in R.DOMINIOS],
              anchos=[2, 14], tam=9, pie="Dominios funcionales.")

    doc.h(2, "5.2 Atributos de cada requisito")
    doc.lista([
        ("Criticidad.", "Severidad de la condición de fallo asociada según el capítulo 4. No es una "
         "prioridad de desarrollo."),
        ("Fuente.", "Documento aplicable, objetivo de seguridad o requisito derivado del que procede."),
        ("Método.", "I inspección, A análisis, D demostración, T prueba."),
        ("Estado.", "Situación real a la fecha de emisión. Las tablas emplean las abreviaturas "
         "Verificado (implementado y verificado con evidencia registrada), Impl. / v. pend. "
         "(implementado, verificación formal pendiente), Parcial y Planificado, cuyo significado "
         "completo se detalla en el apartado 9.5."),
    ])

    doc.h(2, "5.3 Especificación por dominio")
    for codigo, nombre in R.DOMINIOS:
        grupo = [h for h in R.HLR if h[0].startswith(f"HLR-{codigo}")]
        if not grupo:
            continue
        doc.h(3, f"5.3.{[c for c, _ in R.DOMINIOS].index(codigo) + 1} {codigo} — {nombre}")
        doc.tabla(["ID", "Requisito", "Crit.", "Fuente", "Mét.", "Estado"],
                  [(rid.replace("HLR-", ""), f"{tit}. {txt}", crit, fuente, met, R.ESTADOS_CORTO[est])
                   for rid, tit, txt, crit, fuente, met, est in grupo],
                  anchos=[1.6, 6.7, 1.1, 3.0, 1.2, 2.4], tam=7.8,
                  pie=f"Requisitos de alto nivel del dominio {codigo}.")
    doc.salto()


def cap6(doc):
    doc.h(1, "6. Requisitos de bajo nivel")

    doc.h(2, "6.1 Criterio de derivación")
    doc.p("Los requisitos de bajo nivel descomponen cada requisito de alto nivel hasta el grado de "
          "detalle necesario para poder implementarlo y verificarlo sobre un componente concreto de "
          "la arquitectura descrita en 2.3. El criterio aplicado es que un requisito de bajo nivel "
          "debe poder verificarse examinando o ejercitando un único componente, sin necesidad de "
          "integrar el sistema completo.")
    doc.p("Cada requisito de bajo nivel referencia el componente de código que lo realiza. Esa "
          "referencia es parte de la especificación y no una anotación informativa: si el código se "
          "reorganiza, la referencia debe actualizarse y la matriz del capítulo 7 reemitirse.")

    doc.h(2, "6.2 Especificación")
    doc.tabla(["ID", "Requisito", "Origen", "Componente", "Mét.", "Estado"],
              [(rid.replace("LLR-", ""), txt, padre.replace("HLR-", ""), comp, met,
                R.ESTADOS_CORTO[est])
               for rid, padre, txt, comp, met, est in R.LLR],
              anchos=[1.6, 6.0, 1.4, 3.4, 1.2, 2.4], tam=7.6,
              pie="Requisitos de bajo nivel y su asignación a componentes.")
    doc.salto()


def cap7(doc):
    doc.h(1, "7. Matriz de trazabilidad")

    doc.h(2, "7.1 Metodología")
    doc.p("La trazabilidad se mantiene en ambos sentidos. En sentido directo demuestra que toda "
          "exigencia acaba implementada y verificada. En sentido inverso demuestra que no existe "
          "función ni código cuya presencia no responda a un requisito, lo que es tan relevante "
          "como lo anterior: el código no justificado es código cuyo comportamiento nadie ha "
          "analizado.")
    doc.figura("fig03_trazabilidad.png", "Cadena de trazabilidad y sentido de cada verificación.",
               ancho_cm=16.4)

    doc.h(2, "7.2 Trazabilidad normativa a requisitos de alto nivel")
    doc.tabla(["Origen normativo", "Materia", "Requisitos de alto nivel"],
              [(a, b, c.replace("HLR-", "")) for a, b, c in R.NORMATIVA_HLR],
              anchos=[4.5, 4.5, 7], tam=8,
              pie="Trazabilidad de la normativa aplicable a los requisitos de alto nivel.")

    doc.h(2, "7.3 Trazabilidad de alto nivel a bajo nivel y a componente")
    doc.p("La tabla se ordena por requisito de alto nivel. La columna de componente identifica el "
          "elemento del árbol de fuentes que realiza el requisito.")
    filas = []
    for rid, tit, *_ in R.HLR:
        hijos = [l for l in R.LLR if l[1] == rid]
        if not hijos:
            filas.append((rid.replace("HLR-", ""), tit, "— sin descomposición —", "", ""))
            continue
        for i, (lid, _, _, comp, _, est) in enumerate(hijos):
            filas.append((rid.replace("HLR-", "") if i == 0 else "",
                          tit if i == 0 else "",
                          lid.replace("LLR-", ""), comp, R.ESTADOS_CORTO[est]))
    doc.tabla(["HLR", "Título del requisito de alto nivel", "LLR", "Componente", "Estado"],
              filas, anchos=[1.6, 4.7, 1.6, 5.0, 2.1], tam=7.6,
              pie="Trazabilidad de requisitos de alto nivel a bajo nivel y a componente de código.")

    doc.h(2, "7.4 Trazabilidad de requisitos a casos de verificación")
    doc.tabla(["Caso", "HLR", "Título del caso", "Nivel", "Mét.", "Entorno", "Criterio de aceptación"],
              [(v[0].replace("VER-", ""), v[1].replace("HLR-", ""), v[2], v[3], v[4], v[5], v[6])
               for v in R.VER],
              anchos=[1.3, 1.5, 3.3, 1.9, 1.1, 2.5, 4.4], tam=7.4,
              pie="Casos de verificación y su trazabilidad a los requisitos de alto nivel.")

    doc.h(2, "7.5 Análisis de cobertura")
    hlr_ids = {h[0] for h in R.HLR}
    llr_padres = {l[1] for l in R.LLR}
    ver_cubiertos = {v[1] for v in R.VER}
    sin_llr = sorted(hlr_ids - llr_padres)
    sin_ver = sorted(hlr_ids - ver_cubiertos)
    llr_huerfanos = sorted({l[0] for l in R.LLR if l[1] not in hlr_ids})
    ver_huerfanos = sorted({v[0] for v in R.VER if v[1] not in hlr_ids})

    doc.tabla(["Comprobación de cobertura", "Resultado"],
              [("Requisitos de alto nivel definidos", str(len(R.HLR))),
               ("Requisitos de bajo nivel definidos", str(len(R.LLR))),
               ("Casos de verificación definidos", str(len(R.VER))),
               ("Requisitos de alto nivel sin descomposición en bajo nivel",
                ", ".join(x.replace("HLR-", "") for x in sin_llr) if sin_llr else "Ninguno"),
               ("Requisitos de alto nivel sin caso de verificación asociado",
                ", ".join(x.replace("HLR-", "") for x in sin_ver) if sin_ver else "Ninguno"),
               ("Requisitos de bajo nivel sin requisito de alto nivel de origen",
                ", ".join(llr_huerfanos) if llr_huerfanos else "Ninguno"),
               ("Casos de verificación sin requisito asociado",
                ", ".join(ver_huerfanos) if ver_huerfanos else "Ninguno"),
               ("Orígenes normativos trazados", str(len(R.NORMATIVA_HLR))),
               ("Requisitos derivados del análisis de seguridad", str(len(R.DERIVADOS)))],
              anchos=[9, 7], tam=9,
              pie="Análisis de cobertura de la trazabilidad.")
    doc.p("La ausencia de requisitos huérfanos en ambos sentidos es condición necesaria para "
          "declarar completa la matriz, pero no suficiente para declarar verificado el producto: la "
          "completitud de la trazabilidad demuestra que existe un plan de verificación para cada "
          "requisito, no que dicho plan se haya ejecutado. El estado real de ejecución se declara en "
          "9.5 y en el capítulo 13.")
    doc.salto()


def cap8(doc):
    doc.h(1, "8. Ciclo de vida de desarrollo")

    doc.h(2, "8.1 Modelo adoptado")
    doc.p("Se adopta un ciclo de vida en modelo en V. La razón de esta elección, frente a modelos "
          "iterativos, es que el modelo en V hace explícita la correspondencia entre cada actividad "
          "de definición y la actividad de verificación que la comprueba, que es precisamente la "
          "estructura de evidencia que exige [A-7] y que la autoridad debe poder auditar.")
    doc.figura("fig01_modelo_v.png", "Ciclo de vida de desarrollo y correspondencia de verificación.")

    doc.h(2, "8.2 Fases, productos y criterios de salida")
    doc.tabla(["Fase", "Entradas", "Productos", "Criterio de salida"],
              [("F1. Requisitos del sistema y ConOps",
                "Necesidad del explotador, marco normativo, entorno de operación",
                "Capítulos 1 a 4 de este documento",
                "Análisis de peligros completo y nivel de aseguramiento propuesto"),
               ("F2. Requisitos de alto nivel",
                "Productos de F1",
                "Capítulo 5; trazabilidad normativa (7.2)",
                "Todo origen normativo trazado a al menos un requisito; requisitos derivados "
                "remitidos al análisis de seguridad"),
               ("F3. Diseño arquitectónico",
                "Productos de F2",
                "Apartado 2.3; definición de interfaces internas y externas (2.4)",
                "Cada requisito de alto nivel asignado a uno o más componentes"),
               ("F4. Requisitos de bajo nivel y diseño detallado",
                "Productos de F3",
                "Capítulo 6; trazabilidad a componente (7.3)",
                "Sin requisitos de alto nivel sin descomponer y sin requisitos de bajo nivel "
                "huérfanos"),
               ("F5. Implementación",
                "Productos de F4, normas de codificación",
                "Código fuente bajo control de versiones",
                "Revisión por pares registrada y construcción reproducible en ambas plataformas"),
               ("F6. Pruebas unitarias",
                "Código y requisitos de bajo nivel",
                "Casos de nivel unitario del árbol test/",
                "Cada requisito de bajo nivel verificado; cobertura estructural conforme a 9.3"),
               ("F7. Pruebas de integración",
                "Componentes verificados",
                "Casos de nivel de integración",
                "Interfaces internas y externas ejercitadas incluyendo condiciones de error"),
               ("F8. Pruebas de sistema",
                "Software integrado",
                "Casos de nivel de sistema (7.4)",
                "Cada requisito de alto nivel verificado con evidencia registrada"),
               ("F9. Validación operacional",
                "Software verificado, procedimientos del explotador",
                "Informe de aceptación operacional",
                "Aceptación formal del explotador y, en su caso, de la autoridad")],
              anchos=[3.2, 3.6, 3.6, 5.6], tam=7.8,
              pie="Fases del ciclo de vida, productos y criterios de salida.")

    doc.h(2, "8.3 Correspondencia de verificación")
    doc.p("La rama derecha del modelo no comprueba el código contra sí mismo, sino contra el "
          "producto de la fase homóloga de la rama izquierda. Las pruebas unitarias verifican los "
          "requisitos de bajo nivel, no la implementación que el propio autor consideró correcta. "
          "Las pruebas de sistema verifican los requisitos de alto nivel. La validación operacional "
          "comprueba algo distinto de todo lo anterior: que el conjunto de requisitos resuelve "
          "efectivamente la necesidad operacional, que es la única actividad capaz de detectar un "
          "requisito correcto pero equivocado.")

    doc.h(2, "8.4 Gestión de cambios dentro del ciclo")
    doc.p("Todo cambio posterior al establecimiento de una línea base entra por el proceso de "
          "gestión de configuración del capítulo 10 y obliga a recorrer de nuevo el tramo del "
          "modelo afectado. En concreto: la modificación de un requisito de bajo nivel obliga a "
          "reejecutar sus pruebas unitarias y a analizar el efecto sobre las de integración; la "
          "modificación de un requisito de alto nivel obliga además a reejecutar sus pruebas de "
          "sistema. El análisis de efecto se documenta y se conserva junto con la petición de "
          "cambio.")
    doc.salto()


def cap9(doc):
    doc.h(1, "9. Verificación y validación")

    doc.h(2, "9.1 Niveles y métodos")
    doc.tabla(["Nivel", "Objeto de verificación", "Responsable", "Independencia"],
              [("Unitario", "Requisitos de bajo nivel sobre un componente aislado",
                "Ingeniería de software", "No exigida"),
               ("Integración", "Interfaces internas y con sistemas externos",
                "Ingeniería de software", "Revisión independiente de los resultados"),
               ("Sistema", "Requisitos de alto nivel sobre el software integrado",
                "Verificación", "Exigida para requisitos de criticidad crítica"),
               ("Operacional", "Adecuación a la necesidad operacional declarada",
                "Explotador con apoyo de OTECH", "Por definición independiente")],
              anchos=[2.4, 6.4, 3.6, 3.6], tam=8.5,
              pie="Niveles de verificación e independencia exigida.")
    doc.p("Los métodos admitidos son inspección (examen del producto sin ejecutarlo), análisis "
          "(razonamiento o cálculo sobre el producto), demostración (ejercicio observado sin "
          "instrumentación) y prueba (ejecución con criterio de aceptación medible). El método de "
          "prueba es obligatorio para todo requisito de criticidad crítica cuyo cumplimiento pueda "
          "expresarse de forma medible.")

    doc.h(2, "9.2 Entornos de verificación")
    doc.tabla(["Entorno", "Descripción", "Uso"],
              [("Banco de escritorio", "Equipo de referencia Windows con enlace simulado",
                "Pruebas unitarias y de integración; casos que no requieren dinámica de vuelo"),
               ("Enlace simulado (MockLink)", "Generador de tráfico MAVLink integrado en el producto",
                "Casos de segregación multiaeronave, carga de mensajes y condiciones de error"),
               ("SIL", "Simulador de vuelo del controlador conectado por red",
                "Casos de misión, contención y modos de vuelo sin aeronave física"),
               ("HIL", "Controlador de vuelo real con simulación de la dinámica",
                "Casos de sistema con comportamiento representativo del programa embarcado real"),
               ("Radio control Android", "Equipo de referencia Android",
                "Casos de interfaz, legibilidad y desempeño en la plataforma portatil"),
               ("Operación real", "Aeronave y entorno reales bajo autorización vigente",
                "Validación operacional exclusivamente")],
              anchos=[3, 6, 7], tam=8.5,
              pie="Entornos de verificación.")

    doc.h(2, "9.3 Cobertura estructural")
    doc.p("Para las funciones a las que se asigna nivel AL4 se exige cobertura de sentencias sobre "
          "el código que las realiza. Para las funciones que se eleven a AL3 conforme a 3.5 se "
          "exige adicionalmente cobertura de decisiones. El código procedente del software "
          "preexistente que no realiza ningún requisito de este documento queda excluido de la "
          "medida de cobertura, y dicha exclusión se justifica caso por caso en el informe de "
          "cobertura; no se admite la exclusión global por origen.")

    doc.h(2, "9.4 Criterios de aceptación")
    doc.p("Un caso de verificación se considera superado únicamente si su criterio de aceptación, "
          "declarado en la tabla del apartado 7.4, se satisface íntegramente y el resultado queda "
          "registrado con identificación de la versión del software, del entorno, de la fecha y del "
          "ejecutante. Un caso ejecutado sobre una versión distinta de la sometida a aprobación no "
          "constituye evidencia válida.")

    doc.h(2, "9.5 Estado actual de la verificación")
    conteo = {k: 0 for k in R.ESTADOS}
    for h in R.HLR:
        conteo[h[6]] += 1
    conteo_l = {k: 0 for k in R.ESTADOS}
    for l in R.LLR:
        conteo_l[l[5]] += 1
    doc.tabla(["Estado (abreviatura en las tablas)", "Significado",
               "Req. de alto nivel", "Req. de bajo nivel"],
              [("Implementado y verificado (Verificado)",
                "Función implementada y comprobada con evidencia registrada, incluida la verificación "
                "sobre controlador de vuelo real",
                str(conteo["IV"]), str(conteo_l["IV"])),
               ("Implementado / verificación pendiente (Impl. / v. pend.)",
                "Función implementada y operativa, sin ejecución formal registrada del caso de "
                "verificación asociado",
                str(conteo["IP"]), str(conteo_l["IP"])),
               ("Parcial",
                "Función implementada solo en parte o dependiente de una configuración no incluida "
                "en la línea base",
                str(conteo["PA"]), str(conteo_l["PA"])),
               ("Planificado",
                "Requisito especificado, no implementado a la fecha de emisión",
                str(conteo["PL"]), str(conteo_l["PL"]))],
              anchos=[3.4, 7.6, 2.5, 2.5], tam=8.5,
              pie="Estado de implementación y verificación a la fecha de emisión.")
    doc.nota("Las cifras anteriores son el dato relevante de este documento para la autoridad. Solo "
             "la primera fila constituye evidencia de verificación; las tres restantes describen "
             "trabajo especificado y planificado, no acreditado. La lectura completa de esta "
             "situación se ofrece en el capítulo 13.")
    doc.salto()


def cap10a12(doc):
    doc.h(1, "10. Gestión de configuración")
    doc.p("Todo el código fuente, la configuración de construcción, los recursos y este documento "
          "se mantienen bajo control de versiones distribuido. La versión del producto se deriva "
          "automáticamente de la etiqueta del repositorio en el momento de la compilación, de modo "
          "que no es posible producir un artefacto cuya versión no corresponda con un estado "
          "identificable del código.")
    doc.tabla(["Actividad", "Regla aplicada"],
              [("Identificación",
                "Cada elemento de configuración se identifica por su ruta en el repositorio y por el "
                "identificador de confirmación que lo contiene."),
               ("Línea base",
                "Una línea base se establece mediante etiqueta anotada y se acompaña del presente "
                "documento en la revisión correspondiente."),
               ("Control de cambios",
                "Todo cambio se incorpora mediante petición revisada por al menos un ingeniero "
                "distinto del autor, con análisis de efecto sobre requisitos y verificación."),
               ("Trazabilidad de la construcción",
                "Los artefactos distribuibles registran la versión del producto, la cadena de "
                "herramientas y las plataformas de destino."),
               ("Archivo",
                "Se conservan los artefactos, los registros de verificación y este documento durante "
                "el periodo exigido por la normativa aplicable al explotador."),
               ("Gestión de problemas",
                "Los defectos se registran con severidad, requisitos afectados y versión en la que "
                "se detectan y se corrigen.")],
              anchos=[3.6, 12.4], tam=9,
              pie="Reglas de gestión de configuración.")

    doc.h(1, "11. Aseguramiento de la calidad del software")
    doc.p("La función de aseguramiento de calidad es organizativamente independiente de la función "
          "de desarrollo y tiene autoridad para bloquear la emisión de una línea base cuyos "
          "criterios de salida (8.2) no se hayan satisfecho.")
    doc.lista([
        ("Revisión por pares.", "Todo cambio en código que realice un requisito de criticidad "
         "crítica se revisa por un ingeniero distinto del autor, quedando registro de la revisión."),
        ("Normas de codificación.", "Se aplican normas de codificación documentadas y su "
         "cumplimiento se comprueba mediante análisis estático en la construcción."),
        ("Auditorías de conformidad.", "Antes de cada línea base se audita la correspondencia entre "
         "el estado real del código y lo declarado en este documento."),
        ("Registro de no conformidades.", "Toda desviación detectada se registra, se clasifica y se "
         "cierra con evidencia; una línea base no se emite con no conformidades críticas abiertas."),
    ])

    doc.h(1, "12. Seguridad de la información")
    doc.p("Se adopta como referencia el proceso de DO-326A [R-6], acotado al alcance del software "
          "de la estación. Las amenazas consideradas con incidencia sobre la seguridad operacional "
          "son la suplantación de la estación frente a la aeronave, la inyección de mandos en el "
          "enlace, la manipulación de los datos cartográficos o de elevación empleados para la "
          "contención, y la alteración de la evidencia registrada.")
    doc.tabla(["Amenaza", "Efecto potencial", "Contramedida", "Requisitos"],
              [("Inyección de mandos en el enlace C2",
                "Mando no solicitado sobre la aeronave (FC-03)",
                "Firma de mensajes MAVLink 2.0 y descarte de firmas inválidas", "SEC-001"),
               ("Distribución de un artefacto alterado",
                "Ejecución de software no autorizado en el puesto de pilotaje",
                "Firma digital del paquete y verificación por el usuario final", "PLT-003"),
               ("Interceptación de credenciales de servicios externos",
                "Suplantación del operador ante el proveedor UTM",
                "Almacenamiento en el almacén seguro de la plataforma", "SEC-002"),
               ("Suplantación de un servicio externo",
                "Datos de contención o de autorización falsos",
                "Validación estricta de la cadena de certificados TLS", "SEC-003"),
               ("Alteración de la evidencia registrada",
                "Pérdida de valor probatorio del registro de la operación",
                "Formato documentado, volcado periódico y control de acceso del sistema anfitrión",
                "REG-004, REG-005")],
              anchos=[3.6, 4.2, 5.4, 2.8], tam=8.5,
              pie="Amenazas consideradas y contramedidas asociadas.")
    doc.salto()


def cap13a14(doc):
    doc.h(1, "13. Limitaciones y declaración de estado de madurez")
    doc.p("Este capítulo declara sin atenuantes la situación real del producto a la fecha de "
          "emisión. Su propósito es evitar que la extensión y el nivel de detalle del resto del "
          "documento se interpreten como una acreditación que no existe.")

    doc.h(2, "13.1 Naturaleza de este documento")
    doc.p("Este documento es una especificación y un plan. No es un informe de cumplimiento, no es "
          "un certificado y no acredita por sí mismo la conformidad del software con ninguna norma. "
          "La conformidad se acredita mediante la ejecución de los casos de verificación del "
          "apartado 7.4 y la conservación de sus registros, actividad que se encuentra en curso.")

    doc.h(2, "13.2 Limitaciones conocidas")
    doc.tabla(["Ref.", "Limitación", "Efecto"],
              [("L-1", "La mayor parte de los requisitos se encuentra en estado implementado con "
                       "verificación formal pendiente (véase 9.5).",
                "No existe evidencia registrada suficiente para sostener una declaración de "
                "conformidad con AL4."),
               ("L-2", "El software se construye sobre un componente preexistente de código abierto "
                       "cuyo desarrollo original no estuvo gobernado por este plan (2.7).",
                "La evidencia de proceso del componente de origen no existe y no puede "
                "reconstruirse; la estrategia adoptada es la verificación sobre el producto integrado."),
               ("L-3", "No se ha ejecutado medida de cobertura estructural sobre el código que "
                       "realiza los requisitos críticos.",
                "El objetivo de cobertura de 9.3 está especificado pero no acreditado."),
               ("L-4", "La verificación sobre controlador de vuelo real se ha realizado sobre una "
                       "única plataforma (ala fija VTOL, PX4).",
                "Los resultados no son extrapolables sin más a otras configuraciones de aeronave."),
               ("L-5", "La integración con proveedores UTM se encuentra parcialmente implementada y "
                       "no se ha verificado contra un proveedor en producción.",
                "Las operaciones que requieran autorización UTM no están cubiertas por este producto "
                "en su estado actual."),
               ("L-6", "Los requisitos de desempeño del dominio PER no disponen de línea base medida.",
                "Los umbrales declarados son objetivos de diseño, no valores verificados."),
               ("L-7", "El paquete distribuible no se firma digitalmente en el procedimiento actual "
                       "de publicación.",
                "El usuario final no dispone de medio para verificar el origen e integridad del "
                "artefacto que instala."),
               ("L-8", "La calificación ambiental y eléctrica del equipo anfitrión no forma parte de "
                       "este documento.",
                "La aptitud del conjunto estación depende de una calificación de equipo que debe "
                "documentarse por separado.")],
              anchos=[1, 6.5, 8.5], tam=8.5,
              pie="Limitaciones conocidas a la fecha de emisión.")

    doc.h(2, "13.3 Lo que sí está acreditado")
    doc.p("Las funciones declaradas en estado implementado y verificado en el apartado 9.5 se han "
          "comprobado sobre el software integrado, y una parte sustancial de ellas frente a un "
          "controlador de vuelo real en configuración de simulación con equipo en el lazo. "
          "Comprenden, en particular, el establecimiento y la supervisión del enlace, la detección "
          "y notificación de su pérdida, la presentación de la telemetría esencial y de la actitud, "
          "la codificación por severidad de los mensajes de la aeronave, la confirmación obligatoria "
          "de los mandos críticos, la prioridad de las alertas frente al video y a los paneles, y la "
          "construcción y ejecución del producto en las dos plataformas soportadas.")

    doc.h(1, "14. Plan de acción para la puesta en vigencia")
    doc.p("Las siguientes acciones son las necesarias para pasar del estado declarado en el capítulo "
          "13 a una declaración de conformidad sostenible ante la autoridad. Se ordenan por "
          "precedencia técnica, no por facilidad de ejecución.")
    doc.tabla(["Ref.", "Acción", "Cierra"],
              [("P-01", "Establecer la línea base formal del código mediante etiqueta anotada y "
                        "asociarla a la revisión A de este documento.", "Requisito previo a P-02"),
               ("P-02", "Redactar y ejecutar los procedimientos de los 76 casos de verificación del "
                        "apartado 7.4, conservando los registros de resultado.", "L-1"),
               ("P-03", "Instrumentar la construcción para medir cobertura de sentencias sobre el "
                        "código que realiza requisitos críticos y emitir el informe de cobertura con "
                        "las exclusiones justificadas.", "L-3"),
               ("P-04", "Incorporar la firma digital del paquete distribuible al procedimiento de "
                        "publicación en ambas plataformas.", "L-7"),
               ("P-05", "Definir el escenario de referencia de desempeño y medir la línea base de "
                        "refresco, capacidad de proceso, memoria y resistencia de 4 horas.", "L-6"),
               ("P-06", "Completar la validación de plan contra altura máxima y geocerca en la "
                        "estación, con sus casos de verificación.", "Estado parcial de MIS-004"),
               ("P-07", "Activar y verificar la firma de mensajes MAVLink 2.0 extremo a extremo.",
                "Estado parcial de SEC-001"),
               ("P-08", "Repetir el conjunto de casos de nivel de sistema sobre al menos una "
                        "plataforma multirrotor adicional.", "L-4"),
               ("P-09", "Verificar la integración UTM contra un proveedor en entorno de "
                        "preproducción o declarar la función fuera del alcance del producto.", "L-5"),
               ("P-10", "Armonizar el identificador de paquete de Android con la identidad del "
                        "producto y documentar la migración.", "Nota del apartado 1.3"),
               ("P-11", "Someter la asignación de nivel de aseguramiento del apartado 3.5 a la "
                        "aceptación de la autoridad competente.", "Requisito previo a la vigencia"),
               ("P-12", "Emitir la revisión B de este documento incorporando los resultados de P-02 "
                        "a P-09 y actualizando el apartado 9.5.", "Cierre del ciclo")],
              anchos=[1, 11, 4], tam=8.5,
              pie="Plan de acción para la puesta en vigencia.")
    doc.salto()


def anexos(doc):
    doc.h(1, "Anexo A. Acrónimos y definiciones")
    doc.tabla(["Término", "Significado"],
              [("ADS-B", "Automatic Dependent Surveillance — Broadcast. Vigilancia dependiente "
                         "automática por radiodifusión."),
               ("AFAC", "Agencia Federal de Aviación Civil (México)."),
               ("AL", "Assurance Level. Nivel de aseguramiento de software según DO-278A."),
               ("BVLOS", "Beyond Visual Line of Sight. Operación más allá del alcance visual."),
               ("C2", "Command and Control. Enlace de mando y control."),
               ("C3", "Command, Control and Communications."),
               ("ConOps", "Concept of Operations. Concepto de operación."),
               ("EVLOS", "Extended Visual Line of Sight. Alcance visual ampliado."),
               ("FHA", "Functional Hazard Assessment. Evaluación funcional de peligros."),
               ("GNSS", "Global Navigation Satellite System."),
               ("HIL", "Hardware In the Loop. Equipo real en el lazo de simulación."),
               ("HLR", "High-Level Requirement. Requisito de alto nivel."),
               ("HMI", "Human-Machine Interface. Interfaz humano-máquina."),
               ("LLR", "Low-Level Requirement. Requisito de bajo nivel."),
               ("MAVLink", "Protocolo de mensajería para vehículos no tripulados."),
               ("OACI / ICAO", "Organización de Aviación Civil Internacional."),
               ("OSO", "Operational Safety Objective. Objetivo de seguridad operacional (SORA)."),
               ("RPA", "Remotely Piloted Aircraft. Aeronave pilotada a distancia."),
               ("RPAS", "Remotely Piloted Aircraft System. Sistema de aeronave pilotada a distancia."),
               ("RTK", "Real-Time Kinematic. Posicionamiento cinemático en tiempo real."),
               ("SAIL", "Specific Assurance and Integrity Level (SORA)."),
               ("SIL", "Software In the Loop. Simulación sin equipo físico."),
               ("SMS", "Safety Management System. Sistema de gestión de la seguridad operacional."),
               ("SORA", "Specific Operations Risk Assessment."),
               ("STATUSTEXT", "Mensaje MAVLink de texto de estado emitido por la aeronave."),
               ("UAS", "Unmanned Aircraft System. Sistema de aeronave no tripulada."),
               ("USS", "UAS Service Supplier. Proveedor de servicios UTM."),
               ("UTM", "UAS Traffic Management. Gestión del tránsito de aeronaves no tripuladas."),
               ("VLOS", "Visual Line of Sight. Operación dentro del alcance visual."),
               ("VTOL", "Vertical Take-Off and Landing.")],
              anchos=[2.6, 13.4], tam=8.5,
              pie="Acrónimos empleados en el documento.")

    doc.h(1, "Anexo B. Escenario de referencia de verificación")
    doc.p("El escenario de referencia es la configuración única frente a la que se ejecutan los "
          "casos de verificación cuyo criterio de aceptación contiene un valor cuantitativo. Un "
          "resultado obtenido en una configuración distinta no es comparable con la línea base.")
    doc.tabla(["Elemento", "Configuración de referencia"],
              [("Aeronave simulada", "Ala fija VTOL, programa embarcado PX4 v1.14.3, en simulación "
                                     "con equipo en el lazo"),
               ("Enlace", "Serie sobre USB a 921 600 bit/s"),
               ("Plan de vuelo", "24 elementos de misión, incluido un patrón de barrido"),
               ("Geocerca", "Polígono de inclusión de 6 vértices y círculo de exclusión"),
               ("Fuente de video", "RTSP H.264 a 1280x720 y 30 imágenes por segundo"),
               ("Duración de la sesión", "60 minutos para los casos funcionales; 4 horas para el "
                                         "caso de resistencia"),
               ("Cartografía", "Teselas satelitales previamente descargadas para el área de operación"),
               ("Equipo anfitrión", "Según el equipo de referencia declarado en el apartado 2.5")],
              anchos=[4, 12], tam=9,
              pie="Escenario de referencia de verificación.")

    doc.h(1, "Anexo C. Hoja de aceptación")
    doc.p("La firma de esta hoja acredita la revisión del documento en la revisión indicada en la "
          "portada y el conocimiento expreso de las limitaciones declaradas en el capítulo 13.")
    doc.tabla(["Organización", "Nombre y cargo", "Firma", "Fecha"],
              [("OTECH — Ingeniería de Software", "", "", ""),
               ("OTECH — Aseguramiento de Calidad", "", "", ""),
               ("OTECH — Dirección Técnica", "", "", ""),
               ("Explotador RPAS", "", "", ""),
               ("Autoridad aeronáutica (acuse de recepción)", "", "", "")],
              anchos=[5.5, 4.5, 3.5, 2.5], tam=9)
    doc.p("Fin del documento.", cursiva=True, tam=9, color=GRIS,
          alinear=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
def construir():
    doc = Doc()
    portada(doc)

    # La portada queda sin encabezado ni pie mediante different_first_page.
    doc._encabezado_pie(doc.d.sections[0])

    control(doc)
    indices(doc)
    cap1(doc)
    cap2(doc)
    cap3(doc)
    cap4(doc)
    cap5(doc)
    cap6(doc)
    cap7(doc)
    cap8(doc)
    cap9(doc)
    cap10a12(doc)
    cap13a14(doc)
    anexos(doc)

    props = doc.d.core_properties
    props.title = f"{CODIGO} — {TITULO}"
    props.subject = f"{PRODUCTO} — {SUBTITULO}"
    props.author = "OTECH — Ingeniería de Software"
    props.category = "Documentación de software aeronáutico"
    props.comments = ("Borrador para revisión. No constituye una declaración de conformidad; "
                      "véase el capítulo 13.")

    salida = os.path.join(DOCS, f"{CODIGO}_Rev{REVISION}.docx")
    doc.d.save(salida)
    print("DOCX ->", salida)
    return salida


if __name__ == "__main__":
    construir()
